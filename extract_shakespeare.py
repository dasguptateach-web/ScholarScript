import docx
doc = docx.Document(r'C:\Users\81\Desktop\ScholarScript Drop\_Processed\Shakespeare and His Plays.docx')

paragraphs = []
for p in doc.paragraphs:
    t = p.text.strip()
    if t:
        paragraphs.append(t)

table = doc.tables[0]
rows = []
for i, row in enumerate(table.rows):
    cells = [cell.text.strip() for cell in row.cells]
    if i == 0:
        continue
    if any(c for c in cells):
        rows.append(cells)

print(f"Paragraphs: {len(paragraphs)}")
print(f"Table rows: {len(rows)}")

lines = []
lines.append("---")
lines.append('title: "Shakespeare And His Plays"')
lines.append("date: 2019-02-28")
lines.append("type: paper")
lines.append('author: "ScholarScript Team"')
lines.append("tags: [shakespeare, plays, elizabethan, drama, literary-analysis]")
lines.append('summary: "A comprehensive overview of Shakespeare\'s plays, their chronological order, dating controversies, authorship question, and a complete character index."')
lines.append("---")

for p in paragraphs:
    lines.append("")
    lines.append(p)

lines.append("")
lines.append("## Complete List of Shakespeare's Characters")
lines.append("")
lines.append("| Character | Lines | Play |")
lines.append("|---|---|---|")
for row in rows:
    escaped = [c.replace("|", "\\|") for c in row]
    lines.append(f"| {escaped[0]} | {escaped[1]} | {escaped[2]} |")

result = "\n".join(lines)
print(f"Total chars: {len(result)}")
print(f"Total lines: {len(lines)}")

with open(r'C:\Users\81\ScholarScript\content\papers\shakespeare-and-his-plays.md', 'w', encoding='utf-8') as f:
    f.write(result)
print("Written successfully")
