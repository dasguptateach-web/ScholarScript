import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from .models import ContentItem
from .parser import slugify


# Optional import guards
_HAS_DOCX = False
_HAS_PDF = False
_HAS_OCR = False
_HAS_ODT = False
_HAS_RTF = False
_HAS_TEX = False

try:
    from docx import Document
    _HAS_DOCX = True
except ImportError:
    pass

try:
    import pdfplumber
    _HAS_PDF = True
except ImportError:
    try:
        import PyPDF2
        _HAS_PDF = True
    except ImportError:
        pass

try:
    import pytesseract
    from PIL import Image
    import pdf2image
    _HAS_OCR = True
except ImportError:
    pass

try:
    from odf.opendocument import load as odf_load
    from odf.text import P
    _HAS_ODT = True
except ImportError:
    pass

try:
    import striprtf.striprtf as rtf_parser
    _HAS_RTF = True
except ImportError:
    pass

try:
    from tex import parse_latex
    _HAS_TEX = True
except ImportError:
    pass


class IngestionEngine:
    """Handles automatic ingestion of documents from /uploads into /content."""

    def __init__(self, uploads_dir: Path, content_dir: Path):
        self.uploads_dir = uploads_dir
        self.content_dir = content_dir
        self.processed_dir = uploads_dir / "processed"
        self.processed_dir.mkdir(exist_ok=True)
        self.results: List[dict] = []

    def ingest_all(self) -> List[dict]:
        """Process all unprocessed files in the uploads directory."""
        self.results = []
        supported = (".doc", ".docx", ".pdf", ".odt", ".rtf", ".txt", ".tex")

        for f in sorted(self.uploads_dir.iterdir()):
            if f.is_file() and f.suffix.lower() in supported:
                result = self._process_single(f)
                self.results.append(result)

        return self.results

    def _process_single(self, filepath: Path) -> dict:
        result = {
            "file": filepath.name,
            "status": "pending",
            "title": "",
            "type": "paper",
            "output": "",
            "error": "",
        }

        try:
            text, metadata = self._extract_text(filepath)
            if not text or not text.strip():
                result["status"] = "error"
                result["error"] = "No text could be extracted"
                return result

            title = self._derive_title(filepath, metadata)
            content_type = self._detect_type(text)
            tags = self._extract_keywords(text, max_keywords=5)
            paper_url = self._extract_doi_or_link(text)
            date = self._derive_date(filepath, metadata)
            md_body = self._text_to_markdown(text, source_ext=filepath.suffix.lower())

            front_matter = self._build_front_matter(
                title=title,
                date=date,
                tags=tags,
                content_type=content_type,
                paper_url=paper_url,
            )

            if content_type == "creative-writing":
                out_dir = self.content_dir / "creative-writing"
            else:
                out_dir = self.content_dir / "papers"
            out_dir.mkdir(exist_ok=True)

            slug = slugify(title)
            out_path = out_dir / f"{slug}.md"

            with open(out_path, "w", encoding="utf-8") as f:
                f.write(front_matter)
                f.write("\n")
                f.write(md_body)

            # Move original to processed
            shutil.move(str(filepath), str(self.processed_dir / filepath.name))

            result["status"] = "success"
            result["title"] = title
            result["type"] = content_type
            result["output"] = str(out_path)

        except Exception as e:
            result["status"] = "error"
            result["error"] = str(e)

        return result

    def _extract_text(self, filepath: Path) -> Tuple[str, dict]:
        """Extract text and metadata from a document file."""
        ext = filepath.suffix.lower()
        metadata = {}

        if ext == ".txt":
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                from .cleaners import clean_text as ct
                return ct(f.read()), metadata

        elif ext == ".docx" and _HAS_DOCX:
            doc = Document(str(filepath))
            paras = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if not t:
                    paras.append("")
                    continue
                style = p.style.name.lower() if p.style else ""
                if "heading" in style or "title" in style or "subtitle" in style:
                    level = re.search(r'heading\s*(\d+)', style)
                    prefix = "#" * min(int(level.group(1)) if level else 2, 6)
                    paras.append(f"{prefix} {t}")
                else:
                    paras.append(t)
            # Extract text from ALL tables
            table_texts = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(c for c in cells):
                        table_rows.append(" | ".join(cells))
                if table_rows:
                    table_texts.append("")
                    for line in table_rows:
                        table_texts.append(line)
            # Preserve paragraph breaks (double newline = paragraph boundary)
            text = "\n\n".join(paras)
            if table_texts:
                text = text + "\n\n" + "\n".join(table_texts)
            from .cleaners import clean_text as ct
            text = ct(text)
            # Extract metadata
            props = doc.core_properties
            if props:
                metadata["title"] = props.title
                metadata["author"] = props.author
                if props.created:
                    metadata["date"] = props.created.isoformat()[:10]
            return text, metadata

        elif ext == ".pdf":
            text = ""
            from .cleaners import clean_pdf_pages, clean_text as ct
            # Try pdfplumber with page-granularity paragraph preservation
            if not text.strip() and _HAS_PDF:
                try:
                    import pdfplumber
                    with pdfplumber.open(str(filepath)) as pdf:
                        raw_pages = []
                        for page in pdf.pages:
                            raw = page.extract_text() or ""
                            raw_pages.append(raw)
                        raw_pages = clean_pdf_pages(raw_pages)
                        page_texts = []
                        for raw in raw_pages:
                            page_texts.append(self._reflow_pdf_text(raw))
                        text = "\n\n".join(page_texts)
                except Exception:
                    pass
            # Fallback to PyPDF2
            if not text.strip():
                try:
                    import PyPDF2
                    with open(filepath, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        raw_pages = []
                        for page in reader.pages:
                            raw = page.extract_text() or ""
                            raw_pages.append(raw)
                        raw_pages = clean_pdf_pages(raw_pages)
                        page_texts = []
                        for raw in raw_pages:
                            page_texts.append(self._reflow_pdf_text(raw))
                        text = "\n\n".join(page_texts)
                except Exception:
                    pass
            if text.strip():
                return text, metadata

            # OCR fallback for scanned PDFs
            if _HAS_OCR:
                try:
                    from pdf2image import convert_from_path
                    import pytesseract
                    images = convert_from_path(str(filepath))
                    text = "\n\n".join(pytesseract.image_to_string(img) for img in images)
                    from .cleaners import clean_text as ct
                    return ct(text), metadata
                except Exception:
                    pass

            return "", metadata

        elif ext == ".odt" and _HAS_ODT:
            try:
                doc = odf_load(str(filepath))
                texts = [str(node) for node in doc.getElementsByType(P)]
                from .cleaners import clean_text as ct
                return ct("\n".join(texts)), metadata
            except Exception:
                pass

        elif ext == ".rtf" and _HAS_RTF:
            try:
                with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                    text = rtf_parser.rtf_to_text(f.read())
                if text:
                    from .cleaners import clean_text as ct
                    return ct(text), metadata
            except Exception:
                pass

        elif ext == ".tex" and _HAS_TEX:
            try:
                text = parse_latex(filepath.read_text(encoding="utf-8"))
                from .cleaners import clean_text as ct
                return ct(text), metadata
            except Exception:
                # Fallback: extract text between { and } in \section{}, \subsection{}, etc.
                raw = filepath.read_text(encoding="utf-8")
                text = re.sub(r"\\(?:section|subsection|subsubsection|textbf|textit|emph)\{([^}]*)\}", r"\1", raw)
                text = re.sub(r"\\[a-zA-Z]+(\{[^}]*\})?", "", text)
                text = re.sub(r"[{}]", "", text)
                text = re.sub(r"%.*", "", text)
                from .cleaners import clean_text as ct
                return ct(text.strip()), metadata

        elif ext == ".doc":
            # .doc files require antiword or catdoc - try system call
            try:
                import subprocess
                result = subprocess.run(["antiword", str(filepath)], capture_output=True, text=True, timeout=10)
                if result.returncode == 0 and result.stdout.strip():
                    from .cleaners import clean_text as ct
                    return ct(result.stdout), metadata
            except Exception:
                pass
            try:
                import subprocess
                result = subprocess.run(["catdoc", str(filepath)], capture_output=True, text=True, timeout=10)
                if result.returncode == 0 and result.stdout.strip():
                    from .cleaners import clean_text as ct
                    return ct(result.stdout), metadata
            except Exception:
                pass

        return "", metadata

    def _derive_title(self, filepath: Path, metadata: dict) -> str:
        title = metadata.get("title", "")
        if title and len(title) > 2:
            return title.strip()
        name = filepath.stem
        name = re.sub(r"[-_]+", " ", name)
        name = re.sub(r"\s+", " ", name).strip()
        name = name.title()
        return name or "Untitled Document"

    def _derive_date(self, filepath: Path, metadata: dict) -> str:
        date = metadata.get("date", "")
        if date:
            return date
        try:
            mtime = filepath.stat().st_mtime
            return datetime.fromtimestamp(mtime).strftime("%Y-%m-%d")
        except Exception:
            return datetime.now().strftime("%Y-%m-%d")

    def _detect_type(self, text: str) -> str:
        """Auto-detect content type from text analysis."""
        lines = text.strip().split("\n")
        total_words = len(text.split())
        if total_words < 10:
            return "paper"

        # Check for poetic patterns (short lines, rhyming patterns)
        poetic_indicators = 0
        long_prose_lines = 0
        for line in lines[:50]:
            stripped = line.strip()
            if not stripped:
                continue
            word_count = len(stripped.split())
            if word_count <= 8:
                poetic_indicators += 1
            elif word_count > 20:
                long_prose_lines += 1

        sample_lines = min(len(lines[:50]), 50)
        if sample_lines > 0:
            poetry_ratio = poetic_indicators / sample_lines
            prose_ratio = long_prose_lines / sample_lines
            if poetry_ratio > 0.4 and prose_ratio < 0.2:
                return "creative-writing"

        return "paper"

    def _extract_keywords(self, text: str, max_keywords: int = 5) -> list:
        """Simple keyword extraction using YAKE-like approach (TF-based)."""
        import re
        from collections import Counter

        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        stopwords = {
            "the", "and", "for", "are", "but", "not", "you", "all", "can",
            "had", "her", "was", "one", "our", "out", "has", "have", "been",
            "some", "them", "then", "than", "that", "this", "which", "what",
            "when", "where", "with", "will", "their", "there", "would",
            "about", "could", "should", "also", "into", "over", "such",
            "very", "just", "from", "they", "been", "more", "these", "those",
        }
        words = [w for w in words if w not in stopwords and len(w) > 2]
        counter = Counter(words)
        return [w for w, _ in counter.most_common(max_keywords)]

    def _extract_doi_or_link(self, text: str) -> str:
        """Extract DOI, arXiv ID, or persistent URL from text."""
        doi_match = re.search(r'10\.\d{4,}/[-._;()/:A-Za-z0-9]+', text)
        if doi_match:
            doi = doi_match.group(0).rstrip(').,;:')
            return f"https://doi.org/{doi}"

        arxiv_match = re.search(r'arxiv:\s*(\d{4}\.\d+)', text, re.IGNORECASE)
        if arxiv_match:
            return f"https://arxiv.org/abs/{arxiv_match.group(1)}"

        url_match = re.search(r'https?://(?:dx\.)?doi\.org/\S+', text)
        if url_match:
            return url_match.group(0).rstrip(').,;:')

        return ""

    def _reflow_pdf_text(self, text: str) -> str:
        """Reflow PDF-extracted text to reconstruct paragraphs.

        PDF text extraction typically breaks each physical line with \n.
        This heuristic re-joins lines that belong to the same paragraph
        and inserts blank lines between logical paragraphs.
        """
        lines = text.split("\n")
        blocks = []
        current = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current:
                    blocks.append(" ".join(current))
                    current = []
                continue

            # Hyphenated word-break
            if stripped.endswith("-") and len(stripped) > 3:
                current.append(stripped[:-1])
                continue

            is_short = len(stripped) < 45
            is_very_short = len(stripped) < 25
            starts_upper = stripped[0].isupper() if stripped else False
            previous_short = current and all(len(w) < 40 for w in current)

            if not current:
                current = [stripped]
            elif is_very_short and starts_upper:
                # Very short line starting uppercase = heading or standalone label
                blocks.append(" ".join(current))
                current = [stripped]
            elif previous_short and not is_short and starts_upper:
                # Current block is all short lines, next line is long prose
                # -> flush current as heading and start paragraph
                blocks.append(" ".join(current))
                current = [stripped]
            elif is_short and starts_upper:
                # Short uppercase = new heading/subheading
                blocks.append(" ".join(current))
                current = [stripped]
            else:
                # Continuation of current paragraph
                current.append(stripped)

        if current:
            blocks.append(" ".join(current))

        return "\n\n".join(blocks)

    def _text_to_markdown(self, text: str, source_ext: str = "") -> str:
        """Convert extracted text to Markdown, preserving original paragraph structure.

        Handles two cases:
        - Proper paragraph text (each line = one paragraph, e.g. from docx)
        - Reflowable text (lines within a paragraph broken by newlines, e.g. from PDF)
        """
        raw_blocks = re.split(r'\n\n+', text.strip())
        md_blocks = []

        for block in raw_blocks:
            lines = [ln.strip() for ln in block.split("\n") if ln.strip()]
            if not lines:
                continue

            joined = " ".join(lines)
            joined = re.sub(r'\s+', ' ', joined).strip()
            total_words = len(joined.split())
            first_line = lines[0]
            ends_with_period = joined.rstrip().endswith(('.', '!', '?'))

            # Detect heading: short, all-caps, numbered, or standalone title-like line
            is_heading = (
                total_words <= 12
                and first_line == first_line.upper()
                and total_words >= 2
            ) or (
                total_words <= 15
                and re.match(r'^(?:CHAPTER|Chapter|SECTION|Section|\d+[\.\)])\s', first_line, re.IGNORECASE)
            ) or (
                total_words <= 8
                and not ends_with_period
                and total_words >= 1
                and (first_line[0].isupper() if first_line else False)
            ) or (
                total_words <= 12
                and not ends_with_period
                and total_words >= 2
                and all(w[0].isupper() for w in first_line.split() if w)
            )

            if is_heading:
                md_blocks.append(f"## {joined}")
                md_blocks.append("")
                continue

            # Sub-heading: bold lead-in (e.g. "Introduction." at start)
            if total_words <= 20 and ends_with_period and joined.count(' ') <= 6:
                md_blocks.append(f"**{joined}**")
                md_blocks.append("")
                continue

            # Regular paragraph
            md_blocks.append(joined)
            md_blocks.append("")

        return "\n".join(md_blocks)

    def _build_front_matter(self, title: str, date: str, tags: list,
                            content_type: str, paper_url: str = "") -> str:
        lines = ["---"]
        lines.append(f"title: \"{title}\"")
        lines.append(f"date: {date}")
        lines.append(f"type: {content_type}")
        if tags:
            lines.append(f"tags: [{', '.join(tags)}]")
        if paper_url:
            lines.append(f"paper_url: \"{paper_url}\"")
        lines.append("---")
        return "\n".join(lines)
